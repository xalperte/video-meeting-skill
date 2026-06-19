import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

try:
    import pptx  # noqa: F401
    import docx  # noqa: F401
    from docx import Document
    HAVE_DEPS = True
except ImportError:
    HAVE_DEPS = False


@unittest.skipUnless(HAVE_DEPS, "python-pptx + python-docx required")
class ExecutiveTemplate(unittest.TestCase):
    def _source_docx(self, path):
        """A tiny corporate-like docx: header, example body, no Table Grid/List Bullet."""
        from docx import Document
        from docx.oxml.ns import qn
        doc = Document()
        styles_el = doc.styles.element
        for st in list(styles_el.findall(qn("w:style"))):
            n = st.find(qn("w:name"))
            if n is not None and n.get(qn("w:val")) in ("Table Grid", "List Bullet"):
                styles_el.remove(st)
        doc.add_paragraph("FILLER")
        doc.sections[0].header.paragraphs[0].text = "COMPANY NAME\tExecutive Brief"
        doc.save(path)

    def test_executive_assets(self):
        from build_starter_templates import build_all, STARTERS
        self.assertEqual(STARTERS["executive"]["accent"], "1A2238")
        self.assertEqual(STARTERS["executive"]["heading_font"], "Cambria")
        with tempfile.TemporaryDirectory() as d:
            exe = os.path.join(d, "templates", "presentation", "executive")
            os.makedirs(os.path.join(exe, "source"))
            self._source_docx(os.path.join(exe, "source", "Executive_Template.docx"))
            # 1x1 png logo
            import base64
            with open(os.path.join(exe, "logo.png"), "wb") as fh:
                fh.write(base64.b64decode(
                    b"iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk"
                    b"+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg=="))
            build_all(d)

            # template.yaml: navy header, gold accent, Cambria
            with open(os.path.join(exe, "template.yaml")) as fh:
                y = fh.read()
            self.assertIn("1A2238", y)
            self.assertIn("B08D57", y)
            self.assertIn("Cambria", y)

            # report.docx: has Table Grid + List Bullet, keeps header, body cleared
            doc = Document(os.path.join(exe, "report.docx"))
            names = {s.name for s in doc.styles}
            self.assertIn("Table Grid", names)
            self.assertIn("List Bullet", names)
            self.assertNotIn("FILLER", "\n".join(p.text for p in doc.paragraphs))
            self.assertIn("COMPANY NAME", doc.sections[0].header.paragraphs[0].text)

            # slides.pptx: 0 slides, standard layouts present
            from pptx import Presentation
            prs = Presentation(os.path.join(exe, "slides.pptx"))
            self.assertEqual(len(prs.slides._sldIdLst), 0)
            lnames = {l.name for l in prs.slide_layouts}
            self.assertIn("Title Slide", lnames)
            self.assertIn("Title and Content", lnames)


if __name__ == "__main__":
    unittest.main()
