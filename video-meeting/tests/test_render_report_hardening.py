import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

try:
    from docx import Document
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


def _record():
    return {
        "meeting": {"title": "Sprint", "date": "2026-06-18", "type": "grooming"},
        "summary": {"tldr": "All good", "sections": [
            {"category": "Notes", "points": ["alpha", "beta"]}]},
        "action_items": [{"title": "Do X", "type": "explicit"}],
        "decisions": ["Decided A"],
    }


@unittest.skipUnless(HAVE_DOCX, "python-docx required")
class ReportHardening(unittest.TestCase):
    def _make_base_without_styles(self, tdir):
        """A base docx with example body, a header, and NO Table Grid / List Bullet."""
        from docx import Document
        doc = Document()
        # delete Table Grid + List Bullet styles to simulate a corporate file
        styles_el = doc.styles.element
        from docx.oxml.ns import qn
        for st in list(styles_el.findall(qn("w:style"))):
            name_el = st.find(qn("w:name"))
            if name_el is not None and name_el.get(qn("w:val")) in ("Table Grid", "List Bullet"):
                styles_el.remove(st)
        doc.add_paragraph("EXAMPLE FILLER ONE")
        doc.add_paragraph("EXAMPLE FILLER TWO")
        doc.sections[0].header.paragraphs[0].text = "ACME LETTERHEAD"
        doc.save(os.path.join(tdir, "report.docx"))

    def test_renders_without_styles_and_clears_body_keeps_header(self):
        import json
        from render_report import render
        with tempfile.TemporaryDirectory() as d:
            tdir = os.path.join(d, "tpl"); os.mkdir(tdir)
            self._make_base_without_styles(tdir)
            rec = os.path.join(d, "rec.json")
            with open(rec, "w") as fh:
                json.dump(_record(), fh)
            out_pdf = os.path.join(d, "report.pdf")
            render(rec, out_pdf, "soffice", True, tdir)  # must not raise
            docx = os.path.splitext(out_pdf)[0] + ".docx"
            self.assertTrue(os.path.isfile(docx))
            res = Document(docx)
            text = "\n".join(p.text for p in res.paragraphs)
            self.assertIn("Sprint", text)          # generated content present
            self.assertIn("alpha", text)           # bullet fell back, content kept
            self.assertNotIn("EXAMPLE FILLER", text)  # base body cleared
            hdr = res.sections[0].header.paragraphs[0].text
            self.assertIn("ACME LETTERHEAD", hdr)  # header preserved


if __name__ == "__main__":
    unittest.main()
