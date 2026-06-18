import os
import sys
import tempfile
import unittest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

try:
    from docx import Document
    HAVE_DOCX = True
except ImportError:
    HAVE_DOCX = False


@unittest.skipUnless(HAVE_DOCX, "python-docx required")
class ReportTemplate(unittest.TestCase):
    def _record(self):
        return {
            "meeting": {"title": "Sprint", "date": "2026-06-18", "type": "grooming"},
            "summary": {"tldr": "All good", "sections": [
                {"category": "Notes", "points": ["a"]}]},
            "action_items": [{"title": "Do X", "type": "explicit"}],
            "decisions": ["Decided A"],
        }

    def _render(self, d, template_dir):
        import json
        from render_report import render
        rec = os.path.join(d, "rec.json")
        with open(rec, "w") as fh:
            json.dump(self._record(), fh)
        out_pdf = os.path.join(d, "report.pdf")
        # keep_docx=True, libreoffice may be absent -> PDF step warns, docx stays
        render(rec, out_pdf, "soffice", True, template_dir)
        return os.path.splitext(out_pdf)[0] + ".docx"

    def test_renders_from_base_template(self):
        with tempfile.TemporaryDirectory() as d:
            tdir = os.path.join(d, "tpl")
            os.mkdir(tdir)
            Document().save(os.path.join(tdir, "report.docx"))
            docx = self._render(d, tdir)
            self.assertTrue(os.path.isfile(docx))
            text = "\n".join(p.text for p in Document(docx).paragraphs)
            self.assertIn("Sprint", text)

    def test_no_template_dir_uses_default_doc(self):
        with tempfile.TemporaryDirectory() as d:
            docx = self._render(d, None)
            self.assertTrue(os.path.isfile(docx))


if __name__ == "__main__":
    unittest.main()
