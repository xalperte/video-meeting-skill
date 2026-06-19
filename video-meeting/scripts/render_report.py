#!/usr/bin/env python3
"""
Renderer — report.pdf from meeting_record.json (runs in render-env: python-docx).

Same content as the slides, but laid out as a READ-AS-A-DOCUMENT report, not a
deck. We build a .docx (clean, styled, also kept as an editable artifact) and
convert it to PDF with LibreOffice headless.

Usage:
  $RENDER_PY render_report.py --record meeting_record.json \
     --out-pdf report.pdf --libreoffice /usr/bin/soffice
"""
import argparse
import json
import os
import subprocess
import sys

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn  # noqa: E402

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from templates import load_template_yaml  # noqa: E402


def fmt_duration(seconds):
    if not seconds:
        return ""
    s = int(round(seconds))
    h, m = s // 3600, (s % 3600) // 60
    return (f"{h}h {m}m" if h else f"{m}m")


def clear_body(doc):
    """Remove existing body paragraphs and tables, keeping the final sectPr
    (and thus section setup); headers/footers live outside the body."""
    body = doc.element.body
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)


def add_bullet(doc, text):
    """Add a bulleted paragraph, degrading gracefully if styles are absent."""
    have = {s.name for s in doc.styles}
    if "List Bullet" in have:
        return doc.add_paragraph(text, style="List Bullet")
    if "List Paragraph" in have:
        return doc.add_paragraph(text, style="List Paragraph")
    return doc.add_paragraph("• " + text)


def _apply_grid_borders(table):
    """Add thin borders to every cell (fallback when Table Grid is absent)."""
    tbl = table._tbl
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr if cell.tcPr is not None else cell.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is None:
            borders = tcPr.makeelement(qn("w:tcBorders"), {})
            tcPr.append(borders)
        for edge in ("top", "left", "bottom", "right"):
            el = borders.makeelement(qn("w:" + edge), {
                qn("w:val"): "single", qn("w:sz"): "4", qn("w:color"): "999999"})
            borders.append(el)


def styled_table(doc, rows, cols):
    """Add a table styled with Table Grid if available, else manual borders."""
    table = doc.add_table(rows=rows, cols=cols)
    if "Table Grid" in {s.name for s in doc.styles}:
        table.style = "Table Grid"
    else:
        _apply_grid_borders(table)
    return table


def add_items_table(doc, items):
    if not items:
        doc.add_paragraph("(none)")
        return
    table = styled_table(doc, 1, 5)
    for c, head in enumerate(["Title", "Assignee", "Priority", "Source", "Conf."]):
        run = table.rows[0].cells[c].paragraphs[0].add_run(head)
        run.bold = True
    for t in items:
        cells = table.add_row().cells
        cells[0].text = t.get("title", "")
        cells[1].text = t.get("assignee", "") or "—"
        cells[2].text = t.get("priority", "")
        cells[3].text = ", ".join(t.get("source_ts", []) or [])
        cells[4].text = str(t.get("confidence", ""))


def build(doc, record):
    m = record.get("meeting", {})
    doc.add_heading(m.get("title", "Meeting") + " — Meeting Report", level=0)

    meta_bits = []
    if m.get("date"):
        meta_bits.append(m["date"])
    if m.get("type"):
        meta_bits.append(f"Type: {m['type']}")
    if m.get("language_out"):
        meta_bits.append(f"Language: {m['language_out']}")
    if m.get("duration_s"):
        meta_bits.append(f"Duration: {fmt_duration(m['duration_s'])}")
    if meta_bits:
        p = doc.add_paragraph(" · ".join(meta_bits))
        p.runs[0].italic = True

    # Participants
    parts = record.get("participants", [])
    if parts:
        doc.add_heading("Participants", level=1)
        table = styled_table(doc, 1, 3)
        for c, head in enumerate(["Name", "Status", "Confidence"]):
            table.rows[0].cells[c].paragraphs[0].add_run(head).bold = True
        for p in parts:
            cells = table.add_row().cells
            cells[0].text = p.get("name", "") or "(unnamed)"
            cells[1].text = p.get("status", "")
            cells[2].text = str(round(p.get("match_confidence", 0.0), 2))

    # Executive summary
    tldr = record.get("summary", {}).get("tldr", "").strip()
    if tldr:
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(tldr)

    # Summary by category
    sections = record.get("summary", {}).get("sections", [])
    if sections:
        doc.add_heading("Summary", level=1)
        for sec in sections:
            doc.add_heading(sec.get("category", "Notes"), level=2)
            for pt in sec.get("points", []):
                add_bullet(doc, pt)

    # Action items
    items = record.get("action_items", [])
    explicit = [t for t in items if t.get("type") == "explicit"]
    suggested = [t for t in items if t.get("type") == "ai_suggested"]
    doc.add_heading("Action Items", level=1)
    doc.add_heading("Agreed in the meeting", level=2)
    add_items_table(doc, explicit)
    doc.add_heading("Suggested (AI)", level=2)
    add_items_table(doc, suggested)

    # Decisions / open questions
    if record.get("decisions"):
        doc.add_heading("Decisions", level=1)
        for d in record["decisions"]:
            add_bullet(doc, d)
    if record.get("open_questions"):
        doc.add_heading("Open Questions", level=1)
        for q in record["open_questions"]:
            add_bullet(doc, q)


def render(record_path, out_pdf, libreoffice, keep_docx=True, template_dir=None):
    with open(record_path, "r", encoding="utf-8") as fh:
        record = json.load(fh)

    base = os.path.join(template_dir, "report.docx") if template_dir else None
    if base and os.path.isfile(base):
        doc = Document(base)
        clear_body(doc)
        _load_template_yaml_warn_if_missing_styles(doc, template_dir)
    else:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name, style.font.size = "Calibri", Pt(11)

    build(doc, record)

    outdir = os.path.dirname(os.path.abspath(out_pdf)) or "."
    os.makedirs(outdir, exist_ok=True)
    docx_path = os.path.splitext(os.path.abspath(out_pdf))[0] + ".docx"
    doc.save(docx_path)
    sys.stderr.write(f"  report.docx -> {docx_path}\n")

    cmd = [libreoffice, "--headless", "--convert-to", "pdf",
           "--outdir", outdir, docx_path]
    p = subprocess.run(cmd, capture_output=True, text=True, check=False)
    pdf_out = os.path.splitext(docx_path)[0] + ".pdf"
    if p.returncode != 0 or not os.path.isfile(pdf_out):
        sys.stderr.write(f"  warn: PDF conversion failed: {p.stderr.strip()}\n")
        sys.stderr.write("  (the .docx is still available; install LibreOffice for PDF)\n")
    else:
        sys.stderr.write(f"  report.pdf -> {pdf_out}\n")

    if not keep_docx and os.path.isfile(pdf_out):
        os.remove(docx_path)


def _load_template_yaml_warn_if_missing_styles(doc, template_dir):
    """Reserved hook: read template.yaml (for future report options) and warn if
    expected built-in styles are absent from a drop-in document."""
    load_template_yaml(template_dir)  # currently no report-specific keys
    have = {s.name for s in doc.styles}
    for needed in ("Heading 1", "Heading 2", "List Bullet", "Table Grid"):
        if needed not in have:
            sys.stderr.write(
                f"  warn: template report.docx lacks style {needed!r}; "
                "falling back to defaults\n")


def main():
    ap = argparse.ArgumentParser(description="Render a PDF report from the meeting record.")
    ap.add_argument("--record", required=True)
    ap.add_argument("--out-pdf", required=True)
    ap.add_argument("--libreoffice", default="soffice")
    ap.add_argument("--template-dir", default=None)
    ap.add_argument("--no-keep-docx", dest="keep_docx", action="store_false")
    ap.set_defaults(keep_docx=True)
    args = ap.parse_args()

    if not os.path.isfile(args.record):
        sys.exit(f"record not found: {args.record}")
    render(args.record, args.out_pdf, args.libreoffice, args.keep_docx, args.template_dir)


if __name__ == "__main__":
    main()
